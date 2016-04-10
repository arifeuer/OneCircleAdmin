from docx import *
import sys


# REQUIRES IN ARGUEMTNS THE FOLLOWING ORDER (excluding the name of this pythons cript:
# Field Rep
# Certification Number
# Start Date
# End Date
# Location
# Hours Certified
# Date Certified
# Course Title

if len(sys.argv) < 11:
	print("not enough arguments")
	sys.exit()

in_filename = sys.argv[1]
field_rep = sys.argv[2]
cert_number = sys.argv[3]
start_date = sys.argv[4]
end_date = sys.argv[5]
location_to_edit = sys.argv[6]
certified_date = sys.argv[7]
title_course = sys.argv[8]
total_part = sys.argv[9]
out_filename = sys.argv[10]

document = Document(in_filename)


FIELDS = {"FIELD_REP": field_rep, 
			"CERT_NUMBER": cert_number,
			"START_DATE": start_date, 
			"END_DATE": end_date, 
			"LOCATION_TO_EDIT": location_to_edit, 
			"CERTIFIED_DATE": certified_date, 
			"TITLE_COURSE": title_course, 
			"TOTAL_PART": total_part}

for p in document.tables:
	for c in p.columns:
		for n in c.cells:
	
			for field in FIELDS:
				if field in n.text:
					for a in n.paragraphs:
						if field in a.text:
							a.text = a.text.replace(field, FIELDS[field]) 

document.save(out_filename)